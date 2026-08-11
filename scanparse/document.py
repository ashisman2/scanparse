"""Structured document model and exporters (Markdown, JSON, DOCX)."""

from __future__ import annotations

import json
import os
from dataclasses import asdict, dataclass, field
from enum import Enum
from typing import Any


class BlockType(str, Enum):
    """Type of a document block."""

    TITLE = "title"
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    LIST_ITEM = "list_item"
    TABLE = "table"
    CAPTION = "caption"
    FOOTNOTE = "footnote"


@dataclass
class Block:
    """A structural block extracted from a document page."""

    block_type: BlockType
    text: str
    order: int = 0
    language: str | None = None  # "en", "hi", or "mixed"
    confidence: float | None = None
    bbox: list[float] | None = None  # [x0, y0, x1, y1] relative coords


@dataclass
class Page:
    """One page of a parsed document."""

    page_number: int
    blocks: list[Block] = field(default_factory=list)
    language: str | None = None

    @property
    def text(self) -> str:
        return "\n\n".join(b.text for b in sorted(self.blocks, key=lambda b: b.order))


@dataclass
class Document:
    """A fully parsed document with structured output and metadata."""

    source: str
    pages: list[Page] = field(default_factory=list)
    language: list[str] = field(default_factory=lambda: ["en"])
    mode: str = "fast"
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n\n".join(p.text for p in self.pages)

    def to_dict(self) -> dict[str, Any]:
        """Convert the document to a plain dictionary."""
        return {
            "source": self.source,
            "language": self.language,
            "mode": self.mode,
            "metadata": self.metadata,
            "pages": [
                {
                    "page_number": p.page_number,
                    "language": p.language,
                    "blocks": [asdict(b) for b in sorted(p.blocks, key=lambda b: b.order)],
                }
                for p in self.pages
            ],
        }

    def to_json(self, path: str | None = None, indent: int = 2) -> str:
        """Serialize to JSON. If *path* is given, also write to disk."""
        data = json.dumps(self.to_dict(), ensure_ascii=False, indent=indent)
        if path:
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            with open(path, "w", encoding="utf-8") as fh:
                fh.write(data)
        return data

    def to_markdown(self, path: str | None = None) -> str:
        """Render to GitHub-flavored Markdown. If *path* is given, also write to disk."""
        parts = [f"# Document parsed from `{os.path.basename(self.source)}`\n"]
        if self.metadata:
            parts.append(f"> Mode: `{self.mode}` | Languages: {', '.join(self.language)}")
            if "processing_time_s" in self.metadata:
                parts.append(f" | Time: {self.metadata['processing_time_s']:.2f}s")
            parts.append("")

        for page in self.pages:
            if len(self.pages) > 1:
                parts.append(f"<!-- page {page.page_number} -->\n")
            for block in sorted(page.blocks, key=lambda b: b.order):
                lang_tag = (
                    "" if block.language in (None, "mixed", *self.language)
                    else f" [lang: {block.language}]"
                )
                if block.block_type == BlockType.TITLE:
                    parts.append(f"# {block.text}{lang_tag}\n")
                elif block.block_type == BlockType.HEADING:
                    parts.append(f"## {block.text}{lang_tag}\n")
                elif block.block_type == BlockType.CAPTION:
                    parts.append(f"*{block.text}{lang_tag}*\n")
                elif block.block_type == BlockType.FOOTNOTE:
                    parts.append(f"^ [{block.text}]\n")
                elif block.block_type == BlockType.LIST_ITEM:
                    parts.append(f"- {block.text}\n")
                elif block.block_type == BlockType.TABLE:
                    parts.append(_table_to_markdown(block.text))
                    parts.append("")
                else:
                    parts.append(f"{block.text}\n")
        text = "\n".join(parts).strip() + "\n"
        if path:
            os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
            with open(path, "w", encoding="utf-8") as fh:
                fh.write(text)
        return text

    def to_docx(self, path: str) -> str:
        """Render to a .docx file (requires python-docx; falls back to a stub warning)."""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt
        except ImportError:  # pragma: no cover
            raise RuntimeError(
                "python-docx is required for DOCX export. Install with: "
                "pip install 'scanparse[all]' or pip install python-docx"
            )
        doc = DocxDocument()
        for page in self.pages:
            for block in sorted(page.blocks, key=lambda b: b.order):
                style = doc.add_paragraph()
                run = style.add_run(block.text)
                run.font.size = Pt(11)
                if block.block_type == BlockType.TITLE:
                    style.style = doc.styles["Title"]
                elif block.block_type == BlockType.HEADING:
                    style.style = doc.styles["Heading 1"]
                elif block.block_type == BlockType.CAPTION:
                    style.style = doc.styles["Caption"]
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)
        return path


def _table_to_markdown(table_text: str) -> str:
    """Convert a pipe/TSV table string into a Markdown table."""
    rows = [line.strip() for line in table_text.strip().splitlines() if line.strip()]
    if not rows:
        return ""
    def split_row(row: str) -> list[str]:
        if "\t" in row:
            return [c.strip() for c in row.split("\t")]
        if "|" in row:
            return [c.strip() for c in row.strip("|").split("|")]
        return [row]
    ncols = max(len(split_row(r)) for r in rows)
    lines = []
    header = split_row(rows[0]) + [""] * (ncols - len(split_row(rows[0])))
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * ncols) + " |")
    for r in rows[1:]:
        cells = split_row(r) + [""] * (ncols - len(split_row(r)))
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)
